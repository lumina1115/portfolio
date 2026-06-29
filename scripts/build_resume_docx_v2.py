from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path("output/resume")
DOCX_PATH = OUT_DIR / "zongxuezhen_resume.docx"
PHOTO_PATH = Path("tmp/resume_preview/resume_image.png")


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("left", "top", "right", "bottom"):
        edge_data = kwargs.get(edge)
        if not edge_data:
            continue
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_font(run, name="Microsoft YaHei", size=11, bold=False, color="2E3A4D"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, before=0, after=0, line=1.15):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_bullet(cell, label, body):
    paragraph = cell.add_paragraph(style="List Bullet")
    style_paragraph(paragraph, after=1, line=1.02)
    label_run = paragraph.add_run(f"{label}：")
    set_font(label_run, size=10.0, bold=True, color="2F3C4F")
    body_run = paragraph.add_run(body)
    set_font(body_run, size=10.0, color="3D4652")


def add_section_title(cell, title):
    paragraph = cell.add_paragraph()
    style_paragraph(paragraph, before=0, after=3, line=1.0)
    run = paragraph.add_run(title)
    set_font(run, size=13, bold=True, color="2E3A4D")


def add_divider(cell):
    paragraph = cell.add_paragraph()
    style_paragraph(paragraph, before=0, after=4, line=1.0)
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.first_child_found_in("w:pBdr")
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D9DEE5")
    borders.append(bottom)


def add_entry(cell, org, role, date, bullets):
    header = cell.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    header.autofit = False
    header.columns[0].width = Cm(10.7)
    header.columns[1].width = Cm(3.7)

    left = header.cell(0, 0)
    right = header.cell(0, 1)
    for target in (left, right):
        set_cell_margins(target, top=0, start=0, bottom=0, end=0)
        set_cell_border(
            target,
            top={"val": "nil"},
            bottom={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
        )

    p_left = left.paragraphs[0]
    style_paragraph(p_left, after=0, line=1.0)
    run_org = p_left.add_run(org)
    set_font(run_org, size=11.1, bold=True, color="2D3646")
    p_left.add_run().add_break()
    run_role = p_left.add_run(role)
    set_font(run_role, size=9.9, bold=True, color="4E5A69")

    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p_right, after=0, line=1.0)
    run_date = p_right.add_run(date)
    set_font(run_date, size=10.0, bold=True, color="2D3646")

    for bullet in bullets:
        add_bullet(cell, bullet[0], bullet[1])


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.CONTINUOUS
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(0.72)
    section.bottom_margin = Cm(0.68)
    section.left_margin = Cm(1.18)
    section.right_margin = Cm(1.18)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.8)

    document.add_table(rows=0, cols=1)

    outer = document.add_table(rows=1, cols=2)
    outer.alignment = WD_TABLE_ALIGNMENT.CENTER
    outer.autofit = False
    outer.columns[0].width = Cm(13.9)
    outer.columns[1].width = Cm(3.2)

    left = outer.cell(0, 0)
    right = outer.cell(0, 1)
    for target in (left, right):
        set_cell_border(
            target,
            top={"val": "nil"},
            bottom={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
        )
        set_cell_margins(target, top=0, start=0, bottom=0, end=0)

    right.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    p = left.paragraphs[0]
    style_paragraph(p, after=1, line=1.0)
    run = p.add_run("宗雪祯")
    set_font(run, size=20.8, bold=True, color="344253")
    run = p.add_run("   工业设计 / 产品设计")
    set_font(run, size=12.1, color="4C5563")
    contact = left.add_paragraph()
    style_paragraph(contact, after=2, line=1.0)
    for idx, piece in enumerate(("18258943967", "3233507383@qq.com", "杭州")):
        run = contact.add_run(piece)
        set_font(run, size=10.1, color="4A5564")
        if idx < 2:
            spacer = contact.add_run("    ")
            set_font(spacer, size=10.1, color="4A5564")

    if PHOTO_PATH.exists():
        pic = right.paragraphs[0]
        pic.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        style_paragraph(pic, after=0, line=1.0)
        pic.add_run().add_picture(str(PHOTO_PATH), width=Cm(2.45))

    body = document.add_table(rows=1, cols=1)
    body.autofit = False
    body.columns[0].width = Cm(17.4)
    cell = body.cell(0, 0)
    set_cell_border(
        cell,
        top={"val": "nil"},
        bottom={"val": "nil"},
        left={"val": "nil"},
        right={"val": "nil"},
    )
    set_cell_margins(cell, top=0, start=0, bottom=0, end=0)

    add_divider(cell)
    add_section_title(cell, "核心技能")
    add_bullet(cell, "工业设计基础", "具备设计调研、概念推演、产品造型、CMF 表达与基础结构认知能力。")
    add_bullet(cell, "数字表达能力", "熟练使用 Rhino、Keyshot、Photoshop、Illustrator、Figma、Canva 完成建模与视觉呈现。")
    add_bullet(cell, "AIGC 辅助设计", "能够使用 ChatGPT、Claude、Midjourney 辅助概念发散与方案表达提效。")

    add_section_title(cell, "教育背景")
    edu = cell.add_table(rows=1, cols=2)
    edu.alignment = WD_TABLE_ALIGNMENT.CENTER
    edu.autofit = False
    edu.columns[0].width = Cm(10.7)
    edu.columns[1].width = Cm(3.7)
    for target in (edu.cell(0, 0), edu.cell(0, 1)):
        set_cell_margins(target, top=0, start=0, bottom=0, end=0)
        set_cell_border(
            target,
            top={"val": "nil"},
            bottom={"val": "nil"},
            left={"val": "nil"},
            right={"val": "nil"},
        )

    p = edu.cell(0, 0).paragraphs[0]
    style_paragraph(p, after=0, line=1.0)
    run = p.add_run("浙江科技大学")
    set_font(run, size=11.1, bold=True, color="2D3646")
    p.add_run().add_break()
    run = p.add_run("工业设计（本科）")
    set_font(run, size=9.9, bold=True, color="4E5A69")

    p = edu.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, after=0, line=1.0)
    run = p.add_run("2023.09 - 至今")
    set_font(run, size=10.0, bold=True, color="2D3646")
    p.add_run().add_break()
    run = p.add_run("GPA：3.46/5.0（前 30%）")
    set_font(run, size=9.7, bold=True, color="4E5A69")

    p = cell.add_paragraph()
    style_paragraph(p, after=1, line=1.02)
    run = p.add_run("主修课程：")
    set_font(run, size=9.6, bold=True, color="4E5A69")
    run = p.add_run("版式设计、产品摄影、三维建模应用、人机工程学、设计方法学。")
    set_font(run, size=9.6, color="4A5564")

    p = cell.add_paragraph()
    style_paragraph(p, after=3, line=1.02)
    run = p.add_run("代表荣誉：")
    set_font(run, size=9.6, bold=True, color="4E5A69")
    run = p.add_run("省赛一等奖、全国大学生英语竞赛三等奖、校级奖学金，英语六级（CET-6）。")
    set_font(run, size=9.6, color="4A5564")

    add_section_title(cell, "实习经历")
    add_entry(
        cell,
        "浩润建材商行公司",
        "设计项目实习生",
        "2024.07 - 2024.09",
        [
            ("市场调研", "参与洗烘衣机抬高底座项目调研，梳理 30+ 品牌主流尺寸与配色信息，为功能方向定义提供依据。"),
            ("方案迭代", "协同推进 3 版模型优化，结合使用场景与视觉需求调整方案，提升表达完整度与可行性。"),
            ("落地配合", "跟进从概念到模型交付的推进过程，熟悉钣金、塑料等基础工艺，并配合控制外观品质与成本。"),
        ],
    )

    add_section_title(cell, "设计项目经历")
    add_entry(
        cell,
        "映月水晶月饼模具（省赛一等奖）",
        "核心成员",
        "2025.06 - 2025.10",
        [
            ("视觉呈现", "参与 Rhino 建模与 Keyshot 渲染，完成具有水晶质感与节庆氛围的方案表达，提升项目整体呈现度。"),
            ("展板设计", "负责参赛视觉排版，将文化底蕴与产品功能逻辑清晰可视化，支撑竞赛展示与传播表达。"),
            ("项目成果", "项目方案已完成落地上架，具备从概念表达走向实际转化的项目经验。"),
        ],
    )

    add_section_title(cell, "校园经历")
    add_entry(
        cell,
        "浙江科技大学 新媒体运营中心",
        "采编部负责人 / 运营组长",
        "2023.09 - 至今",
        [
            ("团队协作", "负责 5 人团队的选题策划与进度把控，高效对接校内重点活动宣传需求。"),
            ("视觉统筹", "利用 Canva 与秀米重构模板，单篇最高阅读量 10,000+，提升账号内容呈现与传播效率。"),
        ],
    )
    add_entry(
        cell,
        "和山工作室 AIGC 创意营",
        "核心成员 / 设计工作流负责人",
        "2025.05 - 2025.06",
        [],
    )

    document.save(DOCX_PATH)


if __name__ == "__main__":
    main()
